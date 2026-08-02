"""Word-native DOCX creation and high-fidelity template editing.

New documents are generated with explicit design-preset tokens and editable
Word objects. Uploaded templates take a stricter path: EvoFlux inventories and
renders the source, then patches only declared WordprocessingML text nodes in a
copied package. Every unrelated package part must retain its original SHA-256.
"""

from __future__ import annotations

from dataclasses import dataclass, field
import hashlib
import json
import os
from pathlib import Path
import re
import shutil
import subprocess
import tempfile
from typing import Annotated, Any, Literal
import zipfile
from xml.etree import ElementTree as ET

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt, RGBColor
from pydantic import (
    BaseModel,
    ConfigDict,
    Field,
    TypeAdapter,
    field_validator,
    model_validator,
)


W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W14 = "http://schemas.microsoft.com/office/word/2010/wordml"
NS = {"w": W, "w14": W14}
CONTENT_PART_RE = re.compile(
    r"^word/(?:document|header\d+|footer\d+|footnotes|endnotes|comments)\.xml$"
)
PLACEHOLDER_RE = re.compile(
    r"\b(?:lorem ipsum|click here|replace me|tbd|todo)\b|\{\{[^{}]+\}\}", re.IGNORECASE
)


def _hex(value: str) -> RGBColor:
    return RGBColor.from_string(value.removeprefix("#"))


PRESETS: dict[str, dict[str, Any]] = {
    "standard_business_brief": {
        "font": "Calibri",
        "body_align": WD_ALIGN_PARAGRAPH.LEFT,
        "body_after": 6,
        "body_line": 1.10,
        "headings": {
            1: (16, "2E74B5", 16, 8),
            2: (13, "2E74B5", 12, 6),
            3: (12, "1F4D78", 8, 4),
        },
        "list_indent": 0.5,
        "list_hanging": 0.25,
        "list_after": 8,
        "table_fill": "F2F4F7",
    },
    "compact_reference_guide": {
        "font": "Calibri",
        "body_align": WD_ALIGN_PARAGRAPH.LEFT,
        "body_after": 6,
        "body_line": 1.25,
        "headings": {
            1: (16, "2E74B5", 18, 10),
            2: (13, "2E74B5", 14, 7),
            3: (12, "1F4D78", 10, 5),
        },
        "list_indent": 0.375,
        "list_hanging": 0.188,
        "list_after": 4,
        "table_fill": "E8EEF5",
    },
    "narrative_proposal": {
        "font": "Calibri",
        "body_align": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "body_after": 8,
        "body_line": 1.333,
        "headings": {
            1: (16, "2E74B5", 18, 10),
            2: (13, "2E74B5", 12, 6),
            3: (12, "1F4D78", 8, 4),
        },
        "list_indent": 0.375,
        "list_hanging": 0.194,
        "list_after": 4,
        "table_fill": "F4F6F9",
    },
}


class RichRun(BaseModel):
    model_config = ConfigDict(extra="forbid")
    text: str = Field(min_length=1, max_length=8000)
    bold: bool = False
    italic: bool = False
    underline: bool = False
    color: str | None = Field(default=None, pattern=r"^[0-9A-Fa-f]{6}$")
    link: str | None = Field(default=None, max_length=2000)


class ParagraphBlock(BaseModel):
    model_config = ConfigDict(extra="forbid")
    type: Literal["paragraph"] = "paragraph"
    text: str | None = Field(default=None, max_length=20000)
    runs: list[RichRun] = Field(default_factory=list, max_length=200)
    role: Literal["body", "lead", "caption", "source"] = "body"
    align: Literal["left", "center", "right", "justify"] = "left"

    @model_validator(mode="after")
    def validate_text(self) -> ParagraphBlock:
        if bool(self.text) == bool(self.runs):
            raise ValueError("paragraph requires exactly one of text or runs")
        return self


class HeadingBlock(BaseModel):
    model_config = ConfigDict(extra="forbid")
    type: Literal["heading"] = "heading"
    text: str = Field(min_length=1, max_length=400)
    level: Literal[1, 2, 3] = 1


class ListBlock(BaseModel):
    model_config = ConfigDict(extra="forbid")
    type: Literal["list"] = "list"
    items: list[str] = Field(min_length=1, max_length=200)
    ordered: bool = False
    level: int = Field(default=0, ge=0, le=3)


class TableColumn(BaseModel):
    model_config = ConfigDict(extra="forbid")
    label: str = Field(min_length=1, max_length=300)
    width_dxa: int = Field(ge=500, le=9360)
    align: Literal["left", "center", "right"] = "left"


class TableBlock(BaseModel):
    model_config = ConfigDict(extra="forbid")
    type: Literal["table"] = "table"
    columns: list[TableColumn] = Field(min_length=1, max_length=12)
    rows: list[list[str]] = Field(default_factory=list, max_length=1000)
    caption: str | None = Field(default=None, max_length=400)
    source: str | None = Field(default=None, max_length=1000)

    @model_validator(mode="after")
    def validate_table(self) -> TableBlock:
        if sum(column.width_dxa for column in self.columns) != 9360:
            raise ValueError("table column widths must sum to 9360 DXA")
        if any(len(row) != len(self.columns) for row in self.rows):
            raise ValueError("every table row must match the column count")
        return self


class CalloutBlock(BaseModel):
    model_config = ConfigDict(extra="forbid")
    type: Literal["callout"] = "callout"
    label: str = Field(min_length=1, max_length=120)
    text: str = Field(min_length=1, max_length=6000)
    tone: Literal["info", "positive", "warning", "risk"] = "info"


class QuoteBlock(BaseModel):
    model_config = ConfigDict(extra="forbid")
    type: Literal["quote"] = "quote"
    text: str = Field(min_length=1, max_length=6000)
    attribution: str | None = Field(default=None, max_length=400)


class ImageBlock(BaseModel):
    model_config = ConfigDict(extra="forbid")
    type: Literal["image"] = "image"
    path: str = Field(min_length=1, max_length=2000)
    width_inches: float = Field(default=6.5, ge=0.5, le=6.5)
    alt_text: str = Field(min_length=1, max_length=1000)
    caption: str | None = Field(default=None, max_length=400)
    align: Literal["left", "center", "right"] = "center"


class PageBreakBlock(BaseModel):
    model_config = ConfigDict(extra="forbid")
    type: Literal["page_break"] = "page_break"


DocumentBlock = Annotated[
    ParagraphBlock
    | HeadingBlock
    | ListBlock
    | TableBlock
    | CalloutBlock
    | QuoteBlock
    | ImageBlock
    | PageBreakBlock,
    Field(discriminator="type"),
]


class MetadataItem(BaseModel):
    model_config = ConfigDict(extra="forbid")
    label: str = Field(min_length=1, max_length=100)
    value: str = Field(min_length=1, max_length=1000)


class NewDocumentProject(BaseModel):
    model_config = ConfigDict(extra="forbid")
    schema_version: Literal[1] = 1
    mode: Literal["new"] = "new"
    title: str = Field(min_length=1, max_length=500)
    subtitle: str | None = Field(default=None, max_length=1000)
    preset: Literal[
        "standard_business_brief", "compact_reference_guide", "narrative_proposal"
    ]
    header_pattern: Literal[
        "memo_masthead",
        "proposal_centerpiece",
        "editorial_cover",
        "customer_pack",
        "workshop_agenda",
        "customer_story",
    ]
    metadata: list[MetadataItem] = Field(default_factory=list, max_length=30)
    running_header: str | None = Field(default=None, max_length=300)
    footer_label: str | None = Field(default=None, max_length=200)
    blocks: list[DocumentBlock] = Field(min_length=1, max_length=1000)
    author: str = Field(default="EvoFlux", min_length=1, max_length=200)


class ReplaceParagraph(BaseModel):
    model_config = ConfigDict(extra="forbid")
    operation: Literal["replace_paragraph"] = "replace_paragraph"
    part: str = "word/document.xml"
    paragraph: int | None = Field(default=None, ge=0, le=100000)
    para_id: str | None = Field(default=None, pattern=r"^[0-9A-Fa-f]{8}$")
    text: str = Field(max_length=50000)

    @model_validator(mode="after")
    def validate_locator(self) -> ReplaceParagraph:
        if (self.paragraph is None) == (self.para_id is None):
            raise ValueError(
                "replace_paragraph requires exactly one of paragraph or para_id"
            )
        return self


class ReplaceText(BaseModel):
    model_config = ConfigDict(extra="forbid")
    operation: Literal["replace_text"] = "replace_text"
    part: str = "word/document.xml"
    paragraph: int | None = Field(default=None, ge=0, le=100000)
    para_id: str | None = Field(default=None, pattern=r"^[0-9A-Fa-f]{8}$")
    find: str = Field(min_length=1, max_length=10000)
    replace: str = Field(max_length=50000)

    @model_validator(mode="after")
    def validate_locator(self) -> ReplaceText:
        if (self.paragraph is None) == (self.para_id is None):
            raise ValueError(
                "replace_text requires exactly one of paragraph or para_id"
            )
        return self


class ReplaceContentControl(BaseModel):
    model_config = ConfigDict(extra="forbid")
    operation: Literal["replace_content_control"] = "replace_content_control"
    part: str = "word/document.xml"
    tag: str = Field(min_length=1, max_length=500)
    text: str = Field(max_length=50000)


class ReplaceTableCell(BaseModel):
    model_config = ConfigDict(extra="forbid")
    operation: Literal["replace_table_cell"] = "replace_table_cell"
    part: str = "word/document.xml"
    table: int = Field(ge=0, le=10000)
    row: int = Field(ge=0, le=10000)
    column: int = Field(ge=0, le=1000)
    text: str = Field(max_length=50000)


TemplateEdit = Annotated[
    ReplaceParagraph | ReplaceText | ReplaceContentControl | ReplaceTableCell,
    Field(discriminator="operation"),
]


class TemplateDocumentProject(BaseModel):
    model_config = ConfigDict(extra="forbid")
    schema_version: Literal[1] = 1
    mode: Literal["template"] = "template"
    title: str = Field(min_length=1, max_length=500)
    source_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    template_confirmed: Literal[True]
    allow_pagination_change: bool = False
    edits: list[TemplateEdit] = Field(min_length=1, max_length=2000)

    @field_validator("edits")
    @classmethod
    def validate_parts(cls, edits: list[TemplateEdit]) -> list[TemplateEdit]:
        for edit in edits:
            if not CONTENT_PART_RE.fullmatch(edit.part):
                raise ValueError(f"unsupported editable DOCX part: {edit.part}")
        return edits


DocumentProject = NewDocumentProject | TemplateDocumentProject
_PROJECT_ADAPTER = TypeAdapter(Annotated[DocumentProject, Field(discriminator="mode")])


@dataclass
class DocxPipelineResult:
    action: str
    work_dir: Path
    source_docx: Path | None = None
    output: Path | None = None
    manifest_path: Path | None = None
    previews: list[Path] = field(default_factory=list)
    issues: list[dict[str, Any]] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)

    @property
    def passed(self) -> bool:
        return not any(issue.get("severity") == "error" for issue in self.issues)

    def to_dict(self) -> dict[str, Any]:
        return {
            "action": self.action,
            "work_dir": str(self.work_dir),
            "source_docx": str(self.source_docx) if self.source_docx else None,
            "output": str(self.output) if self.output else None,
            "manifest_path": str(self.manifest_path) if self.manifest_path else None,
            "previews": [str(path) for path in self.previews],
            "issues": self.issues,
            "passed": self.passed,
            **self.metadata,
        }


def docx_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _package_hashes(path: Path) -> dict[str, str]:
    with zipfile.ZipFile(path) as package:
        return {
            name: hashlib.sha256(package.read(name)).hexdigest()
            for name in package.namelist()
        }


def load_document_project(path: Path) -> DocumentProject:
    return _PROJECT_ADAPTER.validate_json(path.read_text(encoding="utf-8"))


def document_catalog() -> dict[str, Any]:
    return {
        "workflow": "word-native-docx-with-template-fidelity",
        "presets": list(PRESETS),
        "header_patterns": [
            "memo_masthead",
            "proposal_centerpiece",
            "editorial_cover",
            "customer_pack",
            "workshop_agenda",
            "customer_story",
        ],
        "invariants": [
            "Choose exactly one preset and one first-page header pattern for new documents.",
            "Render every page after each meaningful compose operation.",
            "Treat an uploaded DOCX as immutable and patch only declared stable locators.",
            "Preserve every unrelated ZIP package part byte-for-byte in template mode.",
            "Use real Word paragraphs, styles, numbering, tables, images, headers, footers, and page fields.",
        ],
        "new_project_schema": NewDocumentProject.model_json_schema(),
        "template_project_schema": TemplateDocumentProject.model_json_schema(),
    }


def _find_binary(env_name: str, names: tuple[str, ...]) -> str:
    explicit = os.environ.get(env_name)
    if explicit and Path(explicit).is_file():
        return str(Path(explicit).resolve())
    for name in names:
        if found := shutil.which(name):
            return found
    override = (
        Path.home()
        / ".cache"
        / "codex-runtimes"
        / "codex-primary-runtime"
        / "dependencies"
        / "bin"
        / "override"
    )
    for name in names:
        candidate = override / name
        if candidate.is_file():
            return str(candidate)
    raise RuntimeError(f"Required rendering binary is unavailable: {', '.join(names)}")


def render_docx_pages(
    source: Path, render_dir: Path
) -> tuple[list[Path], list[dict[str, Any]]]:
    render_dir.mkdir(parents=True, exist_ok=True)
    soffice = _find_binary("EVOFLUX_SOFFICE_BIN", ("soffice", "libreoffice"))
    pdftoppm = _find_binary("EVOFLUX_PDFTOPPM_BIN", ("pdftoppm",))
    with tempfile.TemporaryDirectory(
        prefix="evoflux-docx-render-", dir=render_dir
    ) as temp:
        temp_dir = Path(temp)
        profile = temp_dir / "profile"
        profile.mkdir()
        env = os.environ.copy()
        env["HOME"] = str(profile)
        env["TMPDIR"] = str(temp_dir)
        conversion = subprocess.run(
            [
                soffice,
                "--headless",
                f"-env:UserInstallation=file://{profile}",
                "--convert-to",
                "pdf",
                "--outdir",
                str(temp_dir),
                str(source),
            ],
            capture_output=True,
            text=True,
            env=env,
            timeout=180,
            check=False,
        )
        pdf = temp_dir / f"{source.stem}.pdf"
        if conversion.returncode != 0 or not pdf.is_file():
            message = (
                conversion.stderr.strip()
                or conversion.stdout.strip()
                or "LibreOffice did not produce a PDF"
            )
            return [], [
                {"severity": "error", "code": "docx-render-failed", "message": message}
            ]
        prefix = temp_dir / "page"
        raster = subprocess.run(
            [pdftoppm, "-png", "-r", "144", str(pdf), str(prefix)],
            capture_output=True,
            text=True,
            timeout=180,
            check=False,
        )
        pages = sorted(temp_dir.glob("page-*.png"))
        if raster.returncode != 0 or not pages:
            return [], [
                {
                    "severity": "error",
                    "code": "docx-raster-failed",
                    "message": raster.stderr.strip() or "pdftoppm produced no pages",
                }
            ]
        outputs = []
        for index, page in enumerate(pages, start=1):
            destination = render_dir / f"page-{index:03d}.png"
            shutil.copy2(page, destination)
            outputs.append(destination)
        return outputs, []


def _text(node: ET.Element) -> str:
    return "".join(item.text or "" for item in node.findall(".//w:t", NS))


def inspect_docx(source: Path, work_dir: Path) -> DocxPipelineResult:
    previews, issues = render_docx_pages(source, work_dir / "previews")
    hashes = _package_hashes(source)
    with zipfile.ZipFile(source) as package:
        parts = []
        for part in sorted(
            name for name in package.namelist() if CONTENT_PART_RE.fullmatch(name)
        ):
            root = ET.fromstring(package.read(part))
            paragraphs = []
            for index, paragraph in enumerate(root.findall(".//w:p", NS)):
                style = paragraph.find("./w:pPr/w:pStyle", NS)
                paragraphs.append(
                    {
                        "paragraph": index,
                        "para_id": paragraph.get(f"{{{W14}}}paraId"),
                        "style": style.get(f"{{{W}}}val")
                        if style is not None
                        else None,
                        "text": _text(paragraph),
                    }
                )
            controls = []
            for control in root.findall(".//w:sdt", NS):
                tag = control.find("./w:sdtPr/w:tag", NS)
                controls.append(
                    {
                        "tag": tag.get(f"{{{W}}}val") if tag is not None else None,
                        "text": _text(control),
                    }
                )
            tables = []
            for table_index, table in enumerate(root.findall(".//w:tbl", NS)):
                rows = table.findall("./w:tr", NS)
                tables.append(
                    {
                        "table": table_index,
                        "rows": len(rows),
                        "columns": max(
                            (len(row.findall("./w:tc", NS)) for row in rows), default=0
                        ),
                    }
                )
            fields = [
                instruction.text or ""
                for instruction in root.findall(".//w:instrText", NS)
            ]
            parts.append(
                {
                    "part": part,
                    "paragraphs": paragraphs,
                    "content_controls": controls,
                    "tables": tables,
                    "fields": fields,
                }
            )
    manifest = {
        "schemaVersion": 1,
        "sourcePath": str(source),
        "sourceSha256": docx_sha256(source),
        "pageCount": len(previews),
        "packageParts": hashes,
        "contentParts": parts,
    }
    manifest_path = work_dir / "document-manifest.json"
    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    manifest_path.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return DocxPipelineResult(
        "inspect",
        work_dir,
        source_docx=source,
        manifest_path=manifest_path,
        previews=previews,
        issues=issues,
        metadata={"page_count": len(previews), "part_count": len(hashes)},
    )


def _set_run_font(
    run: Any, font: str, size: float | None = None, color: str | None = None
) -> None:
    run.font.name = font
    properties = run._element.get_or_add_rPr().get_or_add_rFonts()
    properties.set(qn("w:ascii"), font)
    properties.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = _hex(color)


def _set_cell_margins(
    cell: Any, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120
) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table: Any, widths: list[int], header_fill: str) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    properties = table._tbl.tblPr
    width = properties.first_child_found_in("w:tblW")
    width.set(qn("w:w"), "9360")
    width.set(qn("w:type"), "dxa")
    layout = properties.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    properties.append(indent)
    borders = properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "D9E1E5")
    grid = table._tbl.tblGrid
    grid_columns = grid.findall(qn("w:gridCol"))
    while len(grid_columns) < len(widths):
        grid.append(OxmlElement("w:gridCol"))
        grid_columns = grid.findall(qn("w:gridCol"))
    for index, grid_column in enumerate(grid_columns):
        if index < len(widths):
            grid_column.set(qn("w:w"), str(widths[index]))
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat = header_properties.find(qn("w:tblHeader"))
    if repeat is None:
        repeat = OxmlElement("w:tblHeader")
        header_properties.append(repeat)
    repeat.set(qn("w:val"), "true")
    for row_index, row in enumerate(table.rows):
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.first_child_found_in("w:tcW")
            cell_width.set(qn("w:w"), str(widths[index]))
            cell_width.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
            if row_index == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), header_fill)
                cell._tc.get_or_add_tcPr().append(shading)


def _page_field(paragraph: Any) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))


def _configure_styles(document: DocumentType, preset: dict[str, Any]) -> None:
    normal = document.styles["Normal"]
    normal.font.name = preset["font"]
    normal.font.size = Pt(11)
    normal_fonts = normal.element.get_or_add_rPr().get_or_add_rFonts()
    normal_fonts.set(qn("w:ascii"), preset["font"])
    normal_fonts.set(qn("w:hAnsi"), preset["font"])
    normal.paragraph_format.alignment = preset["body_align"]
    normal.paragraph_format.space_after = Pt(preset["body_after"])
    normal.paragraph_format.line_spacing = preset["body_line"]
    for level, (size, color, before, after) in preset["headings"].items():
        style = document.styles[f"Heading {level}"]
        style.font.name = preset["font"]
        style_fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
        style_fonts.set(qn("w:ascii"), preset["font"])
        style_fonts.set(qn("w:hAnsi"), preset["font"])
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = _hex(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def _build_first_page(
    document: DocumentType, project: NewDocumentProject, preset: dict[str, Any]
) -> None:
    section = document.sections[0]
    header = section.header.paragraphs[0]
    if project.running_header:
        run = header.add_run(project.running_header)
        _set_run_font(run, preset["font"], 9, "66717C")
    footer = section.footer.paragraphs[0]
    if project.footer_label:
        run = footer.add_run(f"{project.footer_label}  |  ")
        _set_run_font(run, preset["font"], 9, "66717C")
    _page_field(footer)

    pattern = project.header_pattern
    if pattern == "editorial_cover":
        for _ in range(4):
            document.add_paragraph()
    kicker_map = {
        "memo_masthead": "EXECUTIVE BRIEF",
        "proposal_centerpiece": "PROPOSAL",
        "editorial_cover": "REPORT",
        "customer_pack": "CUSTOMER PACK",
        "workshop_agenda": "WORKSHOP AGENDA",
        "customer_story": "CUSTOMER STORY",
    }
    centered = pattern in {"proposal_centerpiece", "editorial_cover", "customer_story"}
    opening_alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    )
    kicker = document.add_paragraph()
    kicker.alignment = opening_alignment
    kicker.paragraph_format.space_after = Pt(8)
    run = kicker.add_run(kicker_map[pattern])
    _set_run_font(run, preset["font"], 10, "7A5A00")
    run.bold = True
    title = document.add_paragraph()
    title.alignment = opening_alignment
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run(project.title)
    _set_run_font(
        run, preset["font"], 28 if pattern != "memo_masthead" else 24, "0B2545"
    )
    run.bold = True
    if project.subtitle:
        subtitle = document.add_paragraph()
        subtitle.alignment = opening_alignment
        subtitle.paragraph_format.space_after = Pt(18)
        run = subtitle.add_run(project.subtitle)
        _set_run_font(run, preset["font"], 13, "66717C")
    if project.metadata:
        if pattern in {"proposal_centerpiece", "customer_pack", "workshop_agenda"}:
            table = document.add_table(rows=(len(project.metadata) + 1) // 2, cols=2)
            widths = [4680, 4680]
            for index, item in enumerate(project.metadata):
                cell = table.cell(index // 2, index % 2)
                paragraph = cell.paragraphs[0]
                label = paragraph.add_run(f"{item.label}: ")
                label.bold = True
                paragraph.add_run(item.value)
            _set_table_geometry(table, widths, "FFFFFF")
        else:
            for item in project.metadata:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(2)
                label = paragraph.add_run(f"{item.label}: ")
                label.bold = True
                paragraph.add_run(item.value)
    document.add_paragraph()


def _alignment(value: str) -> WD_ALIGN_PARAGRAPH:
    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[value]


def _add_hyperlink(paragraph: Any, run_spec: RichRun, font: str) -> None:
    relationship = paragraph.part.relate_to(
        run_spec.link, RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), font)
    fonts.set(qn("w:hAnsi"), font)
    properties.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), run_spec.color or "2E74B5")
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    text = OxmlElement("w:t")
    text.text = run_spec.text
    run.extend((properties, text))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_blocks(
    document: DocumentType,
    project: NewDocumentProject,
    preset: dict[str, Any],
    asset_root: Path,
) -> None:
    for block in project.blocks:
        if isinstance(block, HeadingBlock):
            document.add_paragraph(block.text, style=f"Heading {block.level}")
        elif isinstance(block, ParagraphBlock):
            paragraph = document.add_paragraph()
            paragraph.alignment = _alignment(block.align)
            if block.role == "lead":
                paragraph.paragraph_format.space_after = Pt(10)
            elif block.role in {"caption", "source"}:
                paragraph.paragraph_format.space_before = Pt(4)
                paragraph.paragraph_format.space_after = Pt(4)
            if block.text is not None:
                run = paragraph.add_run(block.text)
                _set_run_font(
                    run,
                    preset["font"],
                    11 if block.role not in {"caption", "source"} else 9,
                    "66717C" if block.role in {"caption", "source"} else "24323D",
                )
            else:
                for item in block.runs:
                    if item.link:
                        _add_hyperlink(paragraph, item, preset["font"])
                        continue
                    run = paragraph.add_run(item.text)
                    _set_run_font(run, preset["font"], 11, item.color or "24323D")
                    run.bold = item.bold
                    run.italic = item.italic
                    run.underline = item.underline
        elif isinstance(block, ListBlock):
            for item in block.items:
                paragraph = document.add_paragraph(
                    item, style="List Number" if block.ordered else "List Bullet"
                )
                paragraph.paragraph_format.left_indent = Inches(
                    preset["list_indent"] + block.level * 0.25
                )
                paragraph.paragraph_format.first_line_indent = Inches(
                    -preset["list_hanging"]
                )
                paragraph.paragraph_format.space_after = Pt(preset["list_after"])
                paragraph.paragraph_format.line_spacing = preset["body_line"]
        elif isinstance(block, TableBlock):
            if block.caption:
                caption = document.add_paragraph(block.caption)
                caption.paragraph_format.space_before = Pt(4)
                caption.paragraph_format.space_after = Pt(4)
                caption.runs[0].bold = True
            table = document.add_table(rows=1, cols=len(block.columns))
            for index, column in enumerate(block.columns):
                paragraph = table.rows[0].cells[index].paragraphs[0]
                paragraph.alignment = _alignment(column.align)
                run = paragraph.add_run(column.label)
                run.bold = True
                _set_run_font(run, preset["font"], 9.5, "0B2545")
            for values in block.rows:
                cells = table.add_row().cells
                for index, value in enumerate(values):
                    paragraph = cells[index].paragraphs[0]
                    paragraph.alignment = _alignment(block.columns[index].align)
                    run = paragraph.add_run(value)
                    _set_run_font(run, preset["font"], 9.5, "24323D")
            _set_table_geometry(
                table,
                [column.width_dxa for column in block.columns],
                preset["table_fill"],
            )
            if block.source:
                paragraph = document.add_paragraph(block.source)
                paragraph.paragraph_format.space_before = Pt(4)
                paragraph.paragraph_format.space_after = Pt(4)
                _set_run_font(paragraph.runs[0], preset["font"], 9, "66717C")
        elif isinstance(block, CalloutBlock):
            colors = {
                "info": ("F4F6F9", "1F3A5F"),
                "positive": ("EAF4EE", "23613C"),
                "warning": ("FFF8E1", "7A5A00"),
                "risk": ("FDECEC", "9B1C1C"),
            }
            fill, accent = colors[block.tone]
            table = document.add_table(rows=1, cols=1)
            paragraph = table.cell(0, 0).paragraphs[0]
            label = paragraph.add_run(f"{block.label}\n")
            label.bold = True
            _set_run_font(label, preset["font"], 10, accent)
            text = paragraph.add_run(block.text)
            _set_run_font(text, preset["font"], 10.5, "24323D")
            _set_table_geometry(table, [9360], fill)
        elif isinstance(block, QuoteBlock):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(f"“{block.text}”")
            run.italic = True
            _set_run_font(run, preset["font"], 13, "46505A")
            if block.attribution:
                attribution = document.add_paragraph(f"— {block.attribution}")
                attribution.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_run_font(attribution.runs[0], preset["font"], 9.5, "7A5A00")
        elif isinstance(block, ImageBlock):
            image = Path(block.path)
            image = image if image.is_absolute() else asset_root / image
            if not image.is_file():
                raise FileNotFoundError(f"document image does not exist: {image}")
            paragraph = document.add_paragraph()
            paragraph.alignment = _alignment(block.align)
            run = paragraph.add_run()
            shape = run.add_picture(str(image), width=Inches(block.width_inches))
            properties = shape._inline.docPr
            properties.set("descr", block.alt_text)
            if block.caption:
                caption = document.add_paragraph(block.caption)
                caption.alignment = _alignment(block.align)
                _set_run_font(caption.runs[0], preset["font"], 9, "66717C")
        else:
            document.add_page_break()


def build_new_document(
    project: NewDocumentProject, output: Path, asset_root: Path
) -> None:
    preset = PRESETS[project.preset]
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    _configure_styles(document, preset)
    _build_first_page(document, project, preset)
    _add_blocks(document, project, preset, asset_root)
    document.core_properties.title = project.title
    document.core_properties.author = project.author
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))


def _paragraph(
    root: ET.Element, paragraph_index: int | None, para_id: str | None
) -> ET.Element:
    paragraphs = root.findall(".//w:p", NS)
    if para_id is not None:
        matches = [
            item
            for item in paragraphs
            if (item.get(f"{{{W14}}}paraId") or "").upper() == para_id.upper()
        ]
        if len(matches) != 1:
            raise ValueError(f"Expected one para_id={para_id}, found {len(matches)}")
        return matches[0]
    assert paragraph_index is not None
    try:
        return paragraphs[paragraph_index]
    except IndexError as exc:
        raise ValueError(f"Paragraph index {paragraph_index} is out of range") from exc


def _set_text(node: ET.Element, value: str) -> None:
    texts = node.findall(".//w:t", NS)
    if not texts:
        raise ValueError("Target contains no editable Word text runs")
    texts[0].text = value
    for text in texts[1:]:
        text.text = ""


def _replace_text(node: ET.Element, find: str, replace: str) -> None:
    texts = node.findall(".//w:t", NS)
    combined = "".join(item.text or "" for item in texts)
    start = combined.find(find)
    if start < 0:
        raise ValueError(f"Target does not contain {find!r}")
    end = start + len(find)
    cursor = 0
    inserted = False
    for item in texts:
        value = item.text or ""
        item_start = cursor
        item_end = cursor + len(value)
        cursor = item_end
        if item_end <= start or item_start >= end:
            continue
        left = value[: max(0, start - item_start)]
        right = value[max(0, end - item_start) :]
        if not inserted:
            item.text = left + replace + right
            inserted = True
        else:
            item.text = right


def _apply_template(
    source: Path, project: TemplateDocumentProject, output: Path
) -> dict[str, Any]:
    allowed = {edit.part for edit in project.edits}
    before = _package_hashes(source)
    with zipfile.ZipFile(source) as package:
        roots = {part: ET.fromstring(package.read(part)) for part in allowed}
        for edit in project.edits:
            root = roots[edit.part]
            if isinstance(edit, ReplaceParagraph):
                _set_text(_paragraph(root, edit.paragraph, edit.para_id), edit.text)
            elif isinstance(edit, ReplaceText):
                _replace_text(
                    _paragraph(root, edit.paragraph, edit.para_id),
                    edit.find,
                    edit.replace,
                )
            elif isinstance(edit, ReplaceContentControl):
                matches = []
                for control in root.findall(".//w:sdt", NS):
                    marker = control.find("./w:sdtPr/w:tag", NS)
                    if marker is not None and marker.get(f"{{{W}}}val") == edit.tag:
                        matches.append(control)
                if len(matches) != 1:
                    raise ValueError(
                        f"Expected one content-control tag={edit.tag!r}, found {len(matches)}"
                    )
                _set_text(matches[0], edit.text)
            else:
                try:
                    table = root.findall(".//w:tbl", NS)[edit.table]
                    row = table.findall("./w:tr", NS)[edit.row]
                    cell = row.findall("./w:tc", NS)[edit.column]
                except IndexError as exc:
                    raise ValueError(
                        "Table, row, or column index is out of range"
                    ) from exc
                _set_text(cell, edit.text)
        replacements = {
            part: ET.tostring(root, encoding="utf-8", xml_declaration=True)
            for part, root in roots.items()
        }
        output.parent.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(output, "w") as destination:
            for info in package.infolist():
                destination.writestr(
                    info, replacements.get(info.filename, package.read(info.filename))
                )
    after = _package_hashes(output)
    errors = []
    for part, digest in before.items():
        if part not in allowed and after.get(part) != digest:
            errors.append(f"unrelated package part changed: {part}")
    if set(before) != set(after):
        errors.append("DOCX package part inventory changed")
    return {
        "allowed_parts": sorted(allowed),
        "preserved_part_count": sum(1 for part in before if part not in allowed),
        "errors": errors,
    }


def _structural_qa(path: Path) -> list[dict[str, Any]]:
    issues = []
    try:
        with zipfile.ZipFile(path) as package:
            bad = package.testzip()
            if bad:
                issues.append(
                    {"severity": "error", "code": "corrupt-package", "message": bad}
                )
            for required in (
                "[Content_Types].xml",
                "word/document.xml",
                "word/styles.xml",
            ):
                if required not in package.namelist():
                    issues.append(
                        {
                            "severity": "error",
                            "code": "missing-part",
                            "message": required,
                        }
                    )
            text = " ".join(
                _text(ET.fromstring(package.read(part)))
                for part in package.namelist()
                if CONTENT_PART_RE.fullmatch(part)
            )
            for match in sorted(set(PLACEHOLDER_RE.findall(text))):
                issues.append(
                    {
                        "severity": "warning",
                        "code": "placeholder",
                        "message": str(match),
                    }
                )
    except (zipfile.BadZipFile, ET.ParseError) as exc:
        issues.append(
            {"severity": "error", "code": "invalid-docx", "message": str(exc)}
        )
    return issues


def validate_document_project(
    project: DocumentProject, source: Path | None = None
) -> dict[str, Any]:
    if isinstance(project, TemplateDocumentProject):
        if source is None:
            raise ValueError("template project requires source_docx")
        if docx_sha256(source) != project.source_sha256:
            raise ValueError("source DOCX changed after inspection; inspect it again")
        return {
            "valid": True,
            "mode": "template",
            "edit_count": len(project.edits),
            "editable_parts": sorted({edit.part for edit in project.edits}),
        }
    if source is not None:
        raise ValueError("new document project must not declare source_docx")
    return {
        "valid": True,
        "mode": "new",
        "preset": project.preset,
        "header_pattern": project.header_pattern,
        "block_count": len(project.blocks),
    }


def render_document_project(
    project_path: Path, source: Path | None, *, asset_root: Path, work_dir: Path
) -> DocxPipelineResult:
    project = load_document_project(project_path)
    validate_document_project(project, source)
    preview_docx = work_dir / "preview.docx"
    lineage: dict[str, Any] = {}
    if isinstance(project, NewDocumentProject):
        build_new_document(project, preview_docx, asset_root)
    else:
        assert source is not None
        lineage = _apply_template(source, project, preview_docx)
    previews, render_issues = render_docx_pages(preview_docx, work_dir / "previews")
    issues = _structural_qa(preview_docx) + render_issues
    reference_page_count: int | None = None
    if (
        isinstance(project, TemplateDocumentProject)
        and not project.allow_pagination_change
    ):
        assert source is not None
        reference_previews, reference_issues = render_docx_pages(
            source, work_dir / "reference-previews"
        )
        reference_page_count = len(reference_previews)
        issues.extend(reference_issues)
        if reference_page_count != len(previews):
            issues.append(
                {
                    "severity": "error",
                    "code": "unexpected-pagination-change",
                    "message": (
                        f"template has {reference_page_count} pages but edited document "
                        f"has {len(previews)}"
                    ),
                }
            )
    issues.extend(
        {"severity": "error", "code": "fidelity", "message": error}
        for error in lineage.get("errors", [])
    )
    return DocxPipelineResult(
        "render",
        work_dir,
        source_docx=source,
        previews=previews,
        issues=issues,
        metadata={
            "page_count": len(previews),
            "reference_page_count": reference_page_count,
            "lineage": lineage,
        },
    )


def compose_document_project(
    project_path: Path,
    source: Path | None,
    output: Path,
    *,
    asset_root: Path,
    work_dir: Path,
) -> DocxPipelineResult:
    project = load_document_project(project_path)
    validate_document_project(project, source)
    lineage: dict[str, Any] = {}
    if isinstance(project, NewDocumentProject):
        build_new_document(project, output, asset_root)
    else:
        assert source is not None
        lineage = _apply_template(source, project, output)
    previews, render_issues = render_docx_pages(output, work_dir / "previews")
    issues = _structural_qa(output) + render_issues
    reference_page_count: int | None = None
    if (
        isinstance(project, TemplateDocumentProject)
        and not project.allow_pagination_change
    ):
        assert source is not None
        reference_previews, reference_issues = render_docx_pages(
            source, work_dir / "reference-previews"
        )
        reference_page_count = len(reference_previews)
        issues.extend(reference_issues)
        if reference_page_count != len(previews):
            issues.append(
                {
                    "severity": "error",
                    "code": "unexpected-pagination-change",
                    "message": (
                        f"template has {reference_page_count} pages but edited document "
                        f"has {len(previews)}"
                    ),
                }
            )
    issues.extend(
        {"severity": "error", "code": "fidelity", "message": error}
        for error in lineage.get("errors", [])
    )
    if issues and any(issue["severity"] == "error" for issue in issues):
        output.unlink(missing_ok=True)
        published = None
    else:
        published = output
    return DocxPipelineResult(
        "compose",
        work_dir,
        source_docx=source,
        output=published,
        previews=previews,
        issues=issues,
        metadata={
            "page_count": len(previews),
            "reference_page_count": reference_page_count,
            "lineage": lineage,
        },
    )


__all__ = [
    "DocxPipelineResult",
    "NewDocumentProject",
    "TemplateDocumentProject",
    "compose_document_project",
    "document_catalog",
    "docx_sha256",
    "inspect_docx",
    "load_document_project",
    "render_document_project",
    "render_docx_pages",
    "validate_document_project",
]
