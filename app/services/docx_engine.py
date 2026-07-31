"""Declarative, native DOCX compiler for EvoOffice."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Annotated, Any, Literal, Mapping

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt, RGBColor
from pydantic import BaseModel, ConfigDict, Field, model_validator

from app.agent.builtin_skills.docx.scripts.qa import inspect_docx
from app.agent.builtin_skills.docx.scripts.stylekit import (
    DocumentProfileName,
    DocumentTheme,
    add_page_number_footer,
    apply_theme,
    declare_content_contract,
    set_keep,
    style_table,
)
from app.services.office_visual_qa_service import render_office_images

DOCX_CAPABILITIES: dict[str, dict[str, str]] = {
    "paragraphs_headings_and_styles": {"create": "full", "edit": "full"},
    "real_bullets_and_numbering": {"create": "full", "edit": "full"},
    "tables_and_fixed_geometry": {"create": "full", "edit": "full"},
    "images_alt_text_and_captions": {"create": "full", "edit": "partial"},
    "headers_footers_page_fields": {"create": "full", "edit": "full"},
    "hyperlinks_and_page_breaks": {"create": "full", "edit": "full"},
    "toc_cross_references_and_sections": {"create": "partial", "edit": "partial"},
    "comments_and_tracked_changes": {
        "create": "template-first",
        "edit": "preserve-only",
    },
    "content_controls_and_embedded_objects": {
        "create": "template-first",
        "edit": "preserve-only",
    },
}


class DocxThemeSpec(BaseModel):
    model_config = ConfigDict(extra="forbid")

    body_font: str = Field(default="Calibri", min_length=1, max_length=80)
    heading_font: str = Field(default="Calibri", min_length=1, max_length=80)
    ink: str = Field(default="24323D", pattern=r"^[0-9A-Fa-f]{6}$")
    muted: str = Field(default="66717C", pattern=r"^[0-9A-Fa-f]{6}$")
    accent: str = Field(default="2F6D68", pattern=r"^[0-9A-Fa-f]{6}$")
    table_border: str = Field(default="B8C4C7", pattern=r"^[0-9A-Fa-f]{6}$")

    def to_native(self) -> DocumentTheme:
        return DocumentTheme(**self.model_dump())


class RichRunSpec(BaseModel):
    model_config = ConfigDict(extra="forbid")

    text: str = Field(min_length=1, max_length=4000)
    bold: bool = False
    italic: bool = False
    underline: bool = False
    color: str | None = Field(default=None, pattern=r"^[0-9A-Fa-f]{6}$")
    link: str | None = Field(default=None, max_length=1000)


class ParagraphBlock(BaseModel):
    model_config = ConfigDict(extra="forbid")

    type: Literal["paragraph"] = "paragraph"
    text: str | None = Field(default=None, max_length=12000)
    runs: list[RichRunSpec] = Field(default_factory=list, max_length=100)
    role: Literal["body", "lead", "caption", "source"] = "body"
    align: Literal["left", "center", "right", "justify"] = "left"

    @model_validator(mode="after")
    def validate_content(self) -> ParagraphBlock:
        if bool(self.text) == bool(self.runs):
            raise ValueError("paragraph requires exactly one of text or runs")
        return self


class HeadingBlock(BaseModel):
    model_config = ConfigDict(extra="forbid")

    type: Literal["heading"] = "heading"
    text: str = Field(min_length=1, max_length=240)
    level: Literal[1, 2, 3] = 1


class ListBlock(BaseModel):
    model_config = ConfigDict(extra="forbid")

    type: Literal["list"] = "list"
    items: list[str] = Field(min_length=1, max_length=100)
    ordered: bool = False


class ChecklistBlock(BaseModel):
    model_config = ConfigDict(extra="forbid")

    type: Literal["checklist"] = "checklist"
    items: list[str] = Field(min_length=1, max_length=100)
    checked: list[int] = Field(default_factory=list, max_length=100)


class CalloutBlock(BaseModel):
    model_config = ConfigDict(extra="forbid")

    type: Literal["callout"] = "callout"
    label: str = Field(min_length=1, max_length=80)
    text: str = Field(min_length=1, max_length=3000)
    tone: Literal["info", "positive", "warning", "risk"] = "info"


class QuoteBlock(BaseModel):
    model_config = ConfigDict(extra="forbid")

    type: Literal["quote"] = "quote"
    text: str = Field(min_length=1, max_length=3000)
    attribution: str | None = Field(default=None, max_length=200)


class TableColumnSpec(BaseModel):
    model_config = ConfigDict(extra="forbid")

    label: str = Field(min_length=1, max_length=120)
    width_dxa: int = Field(ge=500, le=8500)
    align: Literal["left", "center", "right"] = "left"


class TableBlock(BaseModel):
    model_config = ConfigDict(extra="forbid")

    type: Literal["table"] = "table"
    columns: list[TableColumnSpec] = Field(min_length=1, max_length=12)
    rows: list[list[str]] = Field(min_length=1, max_length=500)
    caption: str | None = Field(default=None, max_length=240)

    @model_validator(mode="after")
    def validate_table(self) -> TableBlock:
        if sum(column.width_dxa for column in self.columns) != 9360:
            raise ValueError("table column widths must sum to 9360 DXA")
        for index, row in enumerate(self.rows, start=1):
            if len(row) != len(self.columns):
                raise ValueError(
                    f"table row {index} has {len(row)} cells; expected {len(self.columns)}"
                )
        return self


class ImageBlock(BaseModel):
    model_config = ConfigDict(extra="forbid")

    type: Literal["image"] = "image"
    path: str = Field(min_length=1, max_length=1000)
    width_inches: float = Field(default=6.5, ge=1, le=6.5)
    alt_text: str = Field(min_length=1, max_length=500)
    caption: str | None = Field(default=None, max_length=240)
    align: Literal["left", "center", "right"] = "center"


class PageBreakBlock(BaseModel):
    model_config = ConfigDict(extra="forbid")

    type: Literal["page_break"] = "page_break"


DocumentBlock = Annotated[
    ParagraphBlock
    | HeadingBlock
    | ListBlock
    | ChecklistBlock
    | CalloutBlock
    | QuoteBlock
    | TableBlock
    | ImageBlock
    | PageBreakBlock,
    Field(discriminator="type"),
]


class MetadataItem(BaseModel):
    model_config = ConfigDict(extra="forbid")

    label: str = Field(min_length=1, max_length=80)
    value: str = Field(min_length=1, max_length=500)


class DocumentSpec(BaseModel):
    model_config = ConfigDict(extra="forbid")

    title: str = Field(min_length=1, max_length=240)
    subtitle: str | None = Field(default=None, max_length=500)
    profile: DocumentProfileName = "standard-business"
    header_pattern: Literal["simple", "memo_masthead", "editorial_cover"] = "simple"
    theme: DocxThemeSpec = Field(default_factory=DocxThemeSpec)
    metadata: list[MetadataItem] = Field(default_factory=list, max_length=20)
    running_header: str | None = Field(default=None, max_length=160)
    footer_label: str = Field(default="Page ", max_length=80)
    required_sections: list[str] = Field(default_factory=list, max_length=100)
    blocks: list[DocumentBlock] = Field(min_length=1, max_length=500)
    author: str = Field(default="EvoOffice", min_length=1, max_length=120)


@dataclass(frozen=True)
class DocxBuildResult:
    output: Path
    report: dict[str, Any]
    render: dict[str, Any] | None

    @property
    def passed(self) -> bool:
        return not self.report.get("errors") and not (self.render or {}).get("errors")

    def to_dict(self) -> dict[str, Any]:
        return {
            "output": str(self.output),
            "passed": self.passed,
            "report": self.report,
            "render": self.render,
        }


def document_catalog() -> dict[str, Any]:
    return {
        "profiles": [
            "standard-business",
            "compact-reference",
            "narrative-proposal",
            "operational-sop",
        ],
        "header_patterns": ["simple", "memo_masthead", "editorial_cover"],
        "block_types": [
            "paragraph",
            "heading",
            "list",
            "checklist",
            "callout",
            "quote",
            "table",
            "image",
            "page_break",
        ],
        "capabilities": DOCX_CAPABILITIES,
    }


def _set_run_font(run, theme: DocumentTheme, *, size: float | None = None) -> None:
    run.font.name = theme.body_font
    run._element.get_or_add_rPr().get_or_add_rFonts().set(
        qn("w:ascii"), theme.body_font
    )
    run._element.get_or_add_rPr().get_or_add_rFonts().set(
        qn("w:hAnsi"), theme.body_font
    )
    if size is not None:
        run.font.size = Pt(size)


def _add_hyperlink(paragraph, text: str, url: str, theme: DocumentTheme):
    relationship = paragraph.part.relate_to(
        url,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), theme.accent)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend((color, underline))
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend((properties, text_node))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _alignment(value: str):
    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[value]


def _paragraph_shading(paragraph, fill: str, border: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border)
    borders.append(left)
    properties.extend((shading, borders))


def _add_masthead(document, spec: DocumentSpec, theme: DocumentTheme) -> None:
    if spec.header_pattern == "editorial_cover":
        for _ in range(4):
            document.add_paragraph()
        kicker = document.add_paragraph()
        kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = kicker.add_run("EVOOFFICE REPORT")
        _set_run_font(run, theme, size=10)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(theme.accent)
        title = document.add_paragraph(style="Title")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.add_run(spec.title)
        if spec.subtitle:
            subtitle = document.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.paragraph_format.space_after = Pt(18)
            run = subtitle.add_run(spec.subtitle)
            _set_run_font(run, theme, size=14)
            run.font.color.rgb = RGBColor.from_string(theme.muted)
        for item in spec.metadata:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(f"{item.label}: {item.value}")
            _set_run_font(run, theme, size=9.5)
            run.font.color.rgb = RGBColor.from_string(theme.muted)
        document.add_page_break()
        return

    if spec.header_pattern == "memo_masthead":
        kicker = document.add_paragraph()
        kicker.paragraph_format.space_after = Pt(3)
        run = kicker.add_run("EVOOFFICE BRIEF")
        _set_run_font(run, theme, size=9.5)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(theme.accent)
    title = document.add_paragraph(spec.title, style="Title")
    title.paragraph_format.space_after = Pt(4)
    if spec.subtitle:
        subtitle = document.add_paragraph()
        subtitle.paragraph_format.space_after = Pt(14)
        run = subtitle.add_run(spec.subtitle)
        _set_run_font(run, theme, size=13)
        run.font.color.rgb = RGBColor.from_string(theme.muted)
    for item in spec.metadata:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label = paragraph.add_run(f"{item.label}: ")
        label.bold = True
        _set_run_font(label, theme)
        value = paragraph.add_run(item.value)
        _set_run_font(value, theme)


def _render_paragraph(document, block: ParagraphBlock, theme: DocumentTheme) -> None:
    style = "Caption" if block.role in {"caption", "source"} else None
    paragraph = document.add_paragraph(style=style)
    paragraph.alignment = _alignment(block.align)
    if block.role == "lead":
        paragraph.paragraph_format.space_after = Pt(10)
    if block.text is not None:
        run = paragraph.add_run(block.text)
        _set_run_font(run, theme, size=12 if block.role == "lead" else None)
        run.bold = block.role == "lead"
    else:
        for rich in block.runs:
            if rich.link:
                _add_hyperlink(paragraph, rich.text, rich.link, theme)
                continue
            run = paragraph.add_run(rich.text)
            _set_run_font(run, theme)
            run.bold = rich.bold
            run.italic = rich.italic
            run.underline = rich.underline
            if rich.color:
                run.font.color.rgb = RGBColor.from_string(rich.color)
    set_keep(paragraph)


def _render_callout(document, block: CalloutBlock, theme: DocumentTheme) -> None:
    colors = {
        "info": ("EEF3F2", theme.accent),
        "positive": ("E9F3EA", "356B4B"),
        "warning": ("FFF7E3", "9A6A17"),
        "risk": ("FBECEA", "9B2C2C"),
    }
    fill, border = colors[block.tone]
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(9)
    label = paragraph.add_run(f"{block.label.upper()}  ")
    _set_run_font(label, theme)
    label.bold = True
    label.font.color.rgb = RGBColor.from_string(border)
    body = paragraph.add_run(block.text)
    _set_run_font(body, theme)
    _paragraph_shading(paragraph, fill, border)
    set_keep(paragraph)


def _render_quote(document, block: QuoteBlock, theme: DocumentTheme) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.left_indent = Inches(0.55)
    paragraph.paragraph_format.right_indent = Inches(0.55)
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(f"“{block.text}”")
    _set_run_font(run, theme, size=14)
    run.italic = True
    run.font.color.rgb = RGBColor.from_string(theme.ink)
    if block.attribution:
        attribution = document.add_paragraph()
        attribution.alignment = WD_ALIGN_PARAGRAPH.CENTER
        attribution.paragraph_format.space_after = Pt(12)
        run = attribution.add_run(f"— {block.attribution}")
        _set_run_font(run, theme, size=9.5)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(theme.accent)


def _render_table(
    document,
    block: TableBlock,
    theme: DocumentTheme,
    profile: DocumentProfileName,
) -> None:
    if block.caption:
        caption = document.add_paragraph(block.caption, style="Caption")
        set_keep(caption, next_paragraph=True)
    table = document.add_table(rows=1, cols=len(block.columns))
    for index, column in enumerate(block.columns):
        table.rows[0].cells[index].text = column.label
    for row in block.rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    style_table(
        table,
        widths_dxa=[column.width_dxa for column in block.columns],
        theme=theme,
        profile=profile,
    )
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.paragraphs[0].alignment = _alignment(block.columns[index].align)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def _render_image(
    document, block: ImageBlock, theme: DocumentTheme, asset_root: Path
) -> None:
    path = (asset_root / block.path).resolve()
    if not path.is_file() or not path.is_relative_to(asset_root):
        raise ValueError(f"image path is outside asset root or missing: {block.path}")
    paragraph = document.add_paragraph()
    paragraph.alignment = _alignment(block.align)
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(block.width_inches))
    shape._inline.docPr.set("descr", block.alt_text)
    set_keep(paragraph, next_paragraph=bool(block.caption))
    if block.caption:
        caption = document.add_paragraph(block.caption, style="Caption")
        caption.alignment = _alignment(block.align)
        set_keep(caption)


def build_document(
    specification: DocumentSpec | Mapping[str, Any],
    output: Path,
    *,
    asset_root: Path | None = None,
    render_dir: Path | None = None,
) -> DocxBuildResult:
    """Compile a declarative document into editable DOCX and validate it."""

    spec = (
        specification
        if isinstance(specification, DocumentSpec)
        else DocumentSpec.model_validate(specification)
    )
    output = output.expanduser().resolve()
    if output.suffix.lower() != ".docx":
        raise ValueError("Word output must use the .docx extension")
    output.parent.mkdir(parents=True, exist_ok=True)
    assets = (asset_root or output.parent).expanduser().resolve()
    theme = spec.theme.to_native()

    document = Document()
    apply_theme(document, theme, profile=spec.profile)
    if spec.required_sections:
        declare_content_contract(document, spec.required_sections)
    document.core_properties.title = spec.title
    document.core_properties.subject = "Generated by EvoOffice DOCX Engine"
    document.core_properties.author = spec.author
    document.core_properties.keywords = (
        f"{document.core_properties.keywords}; EvoFlux; EvoOffice; editable Word"
    )
    if spec.running_header:
        for section in document.sections:
            paragraph = section.header.paragraphs[0]
            paragraph.text = spec.running_header
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                _set_run_font(run, theme, size=8.5)
                run.font.color.rgb = RGBColor.from_string(theme.muted)
    add_page_number_footer(document, label=spec.footer_label)
    _add_masthead(document, spec, theme)

    for block in spec.blocks:
        if isinstance(block, ParagraphBlock):
            _render_paragraph(document, block, theme)
        elif isinstance(block, HeadingBlock):
            heading = document.add_heading(block.text, level=block.level)
            set_keep(heading, next_paragraph=True)
        elif isinstance(block, ListBlock):
            style = "List Number" if block.ordered else "List Bullet"
            for item in block.items:
                paragraph = document.add_paragraph(item, style=style)
                set_keep(paragraph)
        elif isinstance(block, ChecklistBlock):
            for index, item in enumerate(block.items):
                prefix = "☒" if index in block.checked else "☐"
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.left_indent = Inches(0.25)
                paragraph.paragraph_format.first_line_indent = Inches(-0.18)
                paragraph.add_run(f"{prefix} {item}")
                set_keep(paragraph)
        elif isinstance(block, CalloutBlock):
            _render_callout(document, block, theme)
        elif isinstance(block, QuoteBlock):
            _render_quote(document, block, theme)
        elif isinstance(block, TableBlock):
            _render_table(document, block, theme, spec.profile)
        elif isinstance(block, ImageBlock):
            _render_image(document, block, theme, assets)
        else:
            document.add_page_break()

    document.save(str(output))
    report = inspect_docx(output)
    render = render_office_images(output, render_dir.resolve()) if render_dir else None
    if render and render.get("status") == "rendered":
        report["errors"].extend(render.get("errors", []))
        report["warnings"].extend(render.get("warnings", []))
    return DocxBuildResult(output=output, report=report, render=render)


def validate_document(
    source: Path,
    *,
    render_dir: Path | None = None,
) -> dict[str, Any]:
    source = source.expanduser().resolve()
    report = inspect_docx(source)
    if render_dir:
        render = render_office_images(source, render_dir.resolve())
        report["render"] = render
        if render.get("status") == "rendered":
            report["errors"].extend(render.get("errors", []))
            report["warnings"].extend(render.get("warnings", []))
    return report


__all__ = [
    "DOCX_CAPABILITIES",
    "DocumentSpec",
    "DocxBuildResult",
    "build_document",
    "document_catalog",
    "validate_document",
]
