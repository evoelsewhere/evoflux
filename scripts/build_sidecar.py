#!/usr/bin/env python3
"""Build a relocatable Python sidecar bundle for the desktop shell.

Layout produced under ``<out>/``::

    sidecar-bundle/
      python/                ← python-build-standalone interpreter
        bin/python3
        lib/python3.12/
      site-packages/         ← EvoFlux + dependencies
        app/                 ← API server package
        fastapi/
        pydantic/
        …
      document-runtime/      ← optional document generation/rendering stack
        manifest.json
        node/
        artifact-tool/
        libreoffice/
        poppler/
        fonts/

The Tauri shell runs a tiny bootstrap that adds
``sidecar-bundle/site-packages`` with ``site.addsitedir()`` so platform
``.pth`` files are processed, then runs
``app/cli/__main__.py serve --handshake --generate-token --parent-pid …``.

We deliberately do NOT use ``uv tool install`` — that produces an
isolated venv with absolute paths inside it, which won't survive being
copied into ``Contents/Resources/``. Instead we:

1. Fetch a python-build-standalone tarball for the target triple via
   ``uv python install --install-dir …``.
2. ``uv pip install --target <site-packages> --python <python-bin>``
   the local project + chosen extras.
3. Strip the ``site-packages/`` of caches, tests, docs.
4. Smoke-test the bundle by invoking ``serve --port 0 --handshake``.

Usage::

    python scripts/build_sidecar.py \\
        --root ./ --out ./desktop/sidecar-bundle \\
        --python-version 3.12 [--extras azure-doc-intel]

CI uses this same script on each runner (macos-26, ubuntu-22.04). The
output is consumed by the Tauri bundler via the ``bundle.resources``
entry in ``tauri.conf.json``.
"""

from __future__ import annotations

import argparse
import os
import platform
import shutil
import signal
import subprocess
import sys
import zipfile
from pathlib import Path

try:
    from scripts.build_document_runtime import (
        DOCUMENT_RUNTIME_SHA256_ENV,
        DOCUMENT_RUNTIME_SOURCE_ENV,
        DocumentRuntimeError,
        normalized_architecture,
        normalized_platform,
        stage_document_runtime,
    )
except ModuleNotFoundError:  # Direct ``python scripts/build_sidecar.py`` execution.
    from build_document_runtime import (  # type: ignore[no-redef]
        DOCUMENT_RUNTIME_SHA256_ENV,
        DOCUMENT_RUNTIME_SOURCE_ENV,
        DocumentRuntimeError,
        normalized_architecture,
        normalized_platform,
        stage_document_runtime,
    )

# Patterns to strip from site-packages to shrink the bundle. Runtime Python
# modules and package metadata must survive; static typing artifacts and native
# debug bundles are build-time inputs and are safe to remove from releases.
STRIP_DIR_NAMES = (
    "__pycache__",
    "tests",
    "test",
    "PyObjCTest",  # PyObjC's compiled self-test suite, not runtime bindings
)
STRIP_GLOBS = (
    "**/*.pyc",
    "**/*.pyo",
    "**/*.pyi",
    "**/py.typed",
    "**/*.pdb",  # MSVC debug symbols
    "**/*.dist-info/RECORD",
    # Heavy localization data we don't need:
    "**/locale/*.mo",
)
STRIP_DIR_SUFFIXES = (".dSYM",)  # Apple native debug-symbol bundles
STRIP_RELATIVE_DIRS = (
    # Keep googleapiclient.discovery_cache itself: discovery.build imports it.
    # The bundled static API descriptions are ~100 MiB and a cache miss already
    # falls back to Google's discovery endpoint. EvoFlux does not call this API.
    "googleapiclient/discovery_cache/documents",
)

IS_WINDOWS = platform.system() == "Windows"


def detect_target_triple() -> str:
    """Return the python-build-standalone triple for the current host."""
    system = platform.system()
    machine = platform.machine().lower()
    if system == "Darwin":
        return (
            "aarch64-apple-darwin"
            if machine in ("arm64", "aarch64")
            else "x86_64-apple-darwin"
        )
    if system == "Linux":
        if machine in ("aarch64", "arm64"):
            return "aarch64-unknown-linux-gnu"
        return "x86_64-unknown-linux-gnu"
    if system == "Windows":
        if machine in ("aarch64", "arm64"):
            return "aarch64-pc-windows-msvc"
        return "x86_64-pc-windows-msvc"
    raise SystemExit(f"unsupported host: {system}/{machine}")


def run(cmd: list[str], **kwargs) -> None:
    print(">>", " ".join(cmd), flush=True)
    subprocess.run(cmd, check=True, **kwargs)


def fetch_python(version: str, out: Path) -> Path:
    """Use uv to download python-build-standalone for ``version``.

    Returns the path to the python executable inside the install dir.
    """
    out.mkdir(parents=True, exist_ok=True)
    # ``uv python install --install-dir`` places one or more directories
    # under ``out``. As of uv 0.5+ the layout is:
    #
    #   <out>/cpython-<version>-<triple>/        ← real install root
    #     bin/python3.12
    #     lib/python3.12/
    #     ...
    #   <out>/cpython-<major>-<triple>           ← *symlink* to the versioned dir
    #
    # We must find the real directory, not the major-version symlink, or
    # ``shutil.move()`` later will move the symlink and leave us with a
    # broken pointer at the destination.
    run(
        [
            "uv",
            "python",
            "install",
            "--install-dir",
            str(out),
            version,
        ]
    )
    binary = _find_python_binary(out, version)
    if binary is None:
        listing = "\n  ".join(sorted(str(p) for p in out.iterdir()))
        raise SystemExit(f"no python binary found under {out}. Contents:\n  {listing}")
    return binary


def _find_python_binary(root: Path, version: str) -> Path | None:
    """Locate the python interpreter inside a uv install root.

    Walks ``root`` looking for the canonical executable name(s) and
    returns the first hit that is a *real file* (not a broken symlink).
    """
    # ``python3.X`` is the canonical name in python-build-standalone;
    # ``python3`` is a symlink to it. Prefer the versioned name so the
    # rest of the script doesn't follow a symlink it then has to rewrite
    # during normalisation.
    names = [f"python{version}", "python3"]
    if IS_WINDOWS:
        # python-build-standalone on Windows ships ``python.exe`` and
        # ``python3.12.exe`` at the install root. Prefer the versioned
        # executable to stay consistent with the Unix path.
        names = [f"python{version}.exe", "python.exe"]
    for name in names:
        for candidate in root.rglob(name):
            # ``is_file()`` follows symlinks — we want both that the
            # symlink resolves *and* that the target exists. ``rglob``
            # already excludes broken symlinks on most platforms, but
            # be defensive.
            try:
                if candidate.is_file():
                    return candidate.resolve()
            except OSError:
                continue
    return None


def normalise_python_dir(install_root: Path, target: Path, python_bin: Path) -> Path:
    """Move uv's install tree to a flat ``target/`` directory.

    ``python_bin`` is the resolved (symlink-free) path to the interpreter
    inside ``install_root``. On Unix the interpreter lives at
    ``<root>/bin/python3.X``, so the install root is ``python_bin.parent.parent``.
    On Windows python-build-standalone places ``python.exe`` directly in the
    install root, so the install root is ``python_bin.parent``.

    After normalisation the layout is::

        <target>/bin/python3.12        (Unix)
        <target>/python3.12.exe        (Windows)
        <target>/lib/python3.12/       (Unix)
        <target>/Lib/                  (Windows)
        ...

    Returns the new path of the python binary inside ``target``.
    """
    if IS_WINDOWS:
        source = python_bin.parent
        expected_bin = source / "python.exe"
    else:
        source = python_bin.parent.parent
        expected_bin = source / "bin" / python_bin.name

    if not expected_bin.is_file():
        raise SystemExit(
            f"resolved install root {source} missing expected binary {expected_bin}"
        )

    if target.exists():
        shutil.rmtree(target)
    target.parent.mkdir(parents=True, exist_ok=True)
    # ``shutil.move`` on a directory works across filesystems by falling
    # back to copy + remove. The source might be inside a directory uv
    # also created symlinks into; that's fine — we only move *this*
    # directory, leaving siblings intact.
    shutil.move(str(source), str(target))

    # Compute the new binary path inside ``target`` and verify.
    if IS_WINDOWS:
        new_bin = target / python_bin.name
        if not new_bin.is_file():
            alt = target / "python.exe"
            if alt.is_file():
                new_bin = alt
    else:
        new_bin = target / "bin" / python_bin.name
        if not new_bin.is_file():
            # Fall back to ``python3`` if the rglob picked the versioned
            # name but only ``python3`` exists at the target.
            alt = target / "bin" / "python3"
            if alt.is_file():
                new_bin = alt
    if not new_bin.is_file():
        raise SystemExit(f"normalisation moved tree but binary not at {new_bin}")
    return new_bin


def install_packages(
    python_bin: Path, project_root: Path, site_packages: Path, extras: list[str]
) -> None:
    """Install the local EvoFlux project + extras into ``site_packages``."""
    site_packages.mkdir(parents=True, exist_ok=True)
    spec = "."
    if extras:
        spec = f".[{','.join(extras)}]"
    # uv pip install --target: PEP 668-safe, no virtualenv needed.
    run(
        [
            "uv",
            "pip",
            "install",
            "--python",
            str(python_bin),
            "--target",
            str(site_packages),
            spec,
        ],
        cwd=project_root,
    )


def strip_bundle(site_packages: Path) -> int:
    """Remove caches/tests/etc. from site-packages. Returns bytes saved."""
    removed = 0
    for pattern in STRIP_GLOBS:
        for p in site_packages.glob(pattern):
            try:
                if p.is_file():
                    removed += p.stat().st_size
                    p.unlink()
            except OSError:
                pass

    directories: set[Path] = set()
    for name in STRIP_DIR_NAMES:
        for p in site_packages.rglob(name):
            if p.is_dir():
                directories.add(p)
    for suffix in STRIP_DIR_SUFFIXES:
        directories.update(p for p in site_packages.rglob(f"*{suffix}") if p.is_dir())
    for relative in STRIP_RELATIVE_DIRS:
        candidate = site_packages / relative
        if candidate.is_dir():
            directories.add(candidate)

    # Deepest-first avoids traversing children after a parent has been removed.
    for p in sorted(directories, key=lambda path: len(path.parts), reverse=True):
        if not p.is_dir():
            continue
        try:
            size = sum(f.stat().st_size for f in p.rglob("*") if f.is_file())
            shutil.rmtree(p, ignore_errors=True)
            removed += size
        except OSError:
            pass
    return removed


def zip_pure_python_packages(site_packages: Path) -> tuple[int, int]:
    """Collapse safe pure-Python packages into one zipimport archive.

    Defender's cold-path cost is dominated by opening thousands of small
    files. Only regular packages containing Python source/stubs are moved;
    namespace packages, application code, native extensions, and packages
    with runtime data remain on disk. A .pth file makes the archive available
    through the same ``site.addsitedir`` bootstrap used by the desktop shell.
    """

    archive_path = site_packages / "evoflux-purelib.zip"
    candidates: list[Path] = []
    file_count = 0
    for child in sorted(site_packages.iterdir()):
        if (
            not child.is_dir()
            or child.name == "app"
            or child.name.endswith((".dist-info", ".data"))
            or not (child / "__init__.py").is_file()
        ):
            continue
        files = [path for path in child.rglob("*") if path.is_file()]
        if not files:
            continue
        if any(path.suffix not in {".py", ".pyi"} for path in files):
            continue
        candidates.append(child)
        file_count += len(files)

    if not candidates:
        return 0, 0

    with zipfile.ZipFile(
        archive_path,
        mode="w",
        compression=zipfile.ZIP_DEFLATED,
        compresslevel=6,
    ) as archive:
        for package in candidates:
            for path in package.rglob("*"):
                if path.is_file():
                    archive.write(path, path.relative_to(site_packages))

    for package in candidates:
        shutil.rmtree(package)
    (site_packages / "evoflux-purelib.pth").write_text(
        f"{archive_path.name}\n",
        encoding="utf-8",
    )
    return len(candidates), file_count


def validate_migration_bundle(python_bin: Path, site_packages: Path) -> None:
    """Validate the bundled head marker and an upgrade from the prior revision."""

    migration_root = site_packages.parent / "_migration_smoke"
    env = {
        **os.environ,
        "APP_ENV": "production",
        "EVOFLUX_DATA_DIR": str(migration_root / "data"),
        "EVOFLUX_CONFIG_DIR": str(migration_root / "config"),
        "EVOFLUX_STATE_DIR": str(migration_root / "state"),
        "EVOFLUX_CACHE_DIR": str(migration_root / "cache"),
        "EVOFLUX_WIKI_DIR": str(migration_root / "wiki"),
        "EVOFLUX_WORKSPACE_DIR": str(migration_root / "workspace"),
    }
    script = """
import site, sys
from pathlib import Path
site.addsitedir(sys.argv[1])
from alembic import command
from alembic.config import Config
from alembic.script import ScriptDirectory
from app.core.schema_version import SCHEMA_HEAD
import app

ini = Path(app.__file__).resolve().parent / "alembic.ini"
cfg = Config(str(ini))
scripts = ScriptDirectory.from_config(cfg)
heads = scripts.get_heads()
if heads != [SCHEMA_HEAD]:
    raise SystemExit(
        f"schema marker mismatch: SCHEMA_HEAD={SCHEMA_HEAD!r}, alembic_heads={heads!r}"
    )
head = scripts.get_revision(SCHEMA_HEAD)
previous = head.down_revision
if not isinstance(previous, str):
    raise SystemExit(f"expected a linear migration before {SCHEMA_HEAD}, got {previous!r}")
command.upgrade(cfg, previous)
command.upgrade(cfg, "head")
print(f"migration smoke: {previous} -> {SCHEMA_HEAD} ok")
"""
    try:
        run(
            [str(python_bin), "-c", script, str(site_packages)],
            env=env,
        )
    finally:
        shutil.rmtree(migration_root, ignore_errors=True)


def validate_document_runtime_bundle(
    python_bin: Path,
    site_packages: Path,
    document_runtime: Path,
) -> None:
    """Execute every bundled document-runtime layer before packaging.

    The manifest/checksum gate proves byte integrity. This smoke adds runtime
    compatibility: Node must import artifact-tool, and LibreOffice + Poppler +
    the private font configuration must render a real Office document.
    """
    smoke_root = site_packages.parent / "_document_runtime_smoke"
    env = {
        **os.environ,
        "EVOFLUX_DOCUMENT_RUNTIME_DIR": str(document_runtime),
        "EVOFLUX_DOCUMENT_RUNTIME_SMOKE_DIR": str(smoke_root),
    }
    script = r"""
import os
from pathlib import Path
import site
import subprocess
import sys
import time

site.addsitedir(sys.argv[1])

from docx import Document
from app.services.office.rendering import render_pages
from app.services.office.runtime import resolve_document_runtime

runtime = resolve_document_runtime()
for binary, version_args in (
    (runtime.node, ["--version"]),
    (runtime.soffice, ["--version"]),
    (runtime.pdftoppm, ["-v"]),
    (runtime.pdfinfo, ["-v"]),
    (runtime.chromium, ["--version"]),
):
    completed = subprocess.run(
        [str(binary), *version_args],
        capture_output=True,
        check=False,
        text=True,
        timeout=30,
    )
    if completed.returncode != 0:
        raise SystemExit(
            f"document runtime executable failed: {binary}\n{completed.stderr}"
        )

artifact_import = (
    "import(process.argv[1]).then("
    "() => process.stdout.write('artifact-tool import ok\\n'),"
    "(error) => { console.error(error); process.exit(1); },"
    ");"
)
artifact_probe = subprocess.run(
    [
        str(runtime.node),
        "--input-type=module",
        "--eval",
        artifact_import,
        runtime.artifact_tool.resolve().as_uri(),
    ],
    capture_output=True,
    check=False,
    text=True,
    timeout=60,
)
if artifact_probe.returncode != 0:
    raise SystemExit(
        "artifact-tool import failed:\n"
        f"{artifact_probe.stderr or artifact_probe.stdout}"
    )

smoke_root = Path(os.environ["EVOFLUX_DOCUMENT_RUNTIME_SMOKE_DIR"])
smoke_root.mkdir(parents=True, exist_ok=True)
html_source = smoke_root / "runtime-smoke.html"
html_source.write_text(
    "<!doctype html><html><body style='margin:0;width:320px;height:180px;"
    "background:#11233f;color:white;font:32px sans-serif;display:grid;"
    "place-items:center'>EvoFlux</body></html>",
    encoding="utf-8",
)
html_preview = smoke_root / "runtime-smoke.png"
chromium_probe = subprocess.Popen(
    [
        str(runtime.chromium),
        "--headless=new",
        "--disable-background-networking",
        "--disable-javascript",
        "--hide-scrollbars",
        "--host-resolver-rules=MAP * 0.0.0.0",
        "--window-size=320,180",
        f"--user-data-dir={smoke_root / 'chromium-profile'}",
        f"--screenshot={html_preview}",
        html_source.resolve().as_uri(),
    ],
    stdout=subprocess.PIPE,
    stderr=subprocess.PIPE,
    text=True,
)
deadline = time.monotonic() + 60
last_size = -1
stable_checks = 0
while time.monotonic() < deadline:
    if chromium_probe.poll() is not None and not html_preview.is_file():
        break
    if html_preview.is_file() and html_preview.stat().st_size > 0:
        size = html_preview.stat().st_size
        stable_checks = stable_checks + 1 if size == last_size else 0
        last_size = size
        if stable_checks >= 2:
            chromium_probe.terminate()
            break
    time.sleep(0.15)
try:
    stdout, stderr = chromium_probe.communicate(timeout=3)
except subprocess.TimeoutExpired:
    chromium_probe.kill()
    stdout, stderr = chromium_probe.communicate(timeout=3)
if not html_preview.is_file() or html_preview.stat().st_size == 0:
    raise SystemExit(
        "Chromium headless render failed:\n"
        f"{stderr or stdout}"
    )
source = smoke_root / "runtime-smoke.docx"
document = Document()
document.add_heading("EvoFlux document runtime", level=1)
document.add_paragraph(
    "Node · artifact-tool · Chromium · LibreOffice · Poppler · fonts"
)
document.save(source)
pages, issues = render_pages(source, smoke_root / "render", code_prefix="runtime")
if issues or len(pages) != 1 or not pages[0].is_file():
    raise SystemExit(f"document runtime render smoke failed: {issues!r}")
print(f"document runtime execution smoke: {runtime.manifest['bundle_version']} ok")
"""
    try:
        run(
            [str(python_bin), "-c", script, str(site_packages)],
            env=env,
        )
    finally:
        shutil.rmtree(smoke_root, ignore_errors=True)


def smoke_test(
    python_bin: Path,
    site_packages: Path,
    document_runtime: Path | None = None,
) -> None:
    """Invoke the sidecar briefly to prove the bundle actually works.

    Stages:

    1. Spawn with the same env vars the desktop shell uses.
    2. Wait for the JSON handshake line on stdout (proves imports + bind).
    3. Hit ``/api/health/live`` (proves lifespan startup).
    4. Verify protected HTTP endpoints reject/accept the desktop token.
    5. Upgrade a real browser-presence WebSocket using the bundled client
       (proves both Uvicorn's transport and the packaged dependency).
    6. Terminate and reap.

    Any failure here fails the build — we never want a broken bundle to
    leave CI.
    """
    import json
    import time
    import urllib.error
    import urllib.parse
    import urllib.request

    smoke_root = site_packages.parent / "_smoke"
    # PYTHONHOME must point at the python-build-standalone install root —
    # the directory containing the standard library. On Unix that's the
    # grandparent of ``bin/python3.X``; on Windows the binary lives at the
    # install root, so it's the parent.
    python_home = python_bin.parent.parent if not IS_WINDOWS else python_bin.parent
    env = {
        **os.environ,
        "PYTHONHOME": str(python_home),
        "PYTHONUNBUFFERED": "1",
        "APP_ENV": "production",
        # Keep test data isolated so the smoke run never touches the user's
        # real EvoFlux directories.
        "EVOFLUX_DATA_DIR": str(smoke_root / "data"),
        "EVOFLUX_CONFIG_DIR": str(smoke_root / "config"),
        "EVOFLUX_STATE_DIR": str(smoke_root / "state"),
        "EVOFLUX_CACHE_DIR": str(smoke_root / "cache"),
        "EVOFLUX_WIKI_DIR": str(smoke_root / "wiki"),
        "EVOFLUX_WORKSPACE_DIR": str(smoke_root / "workspace"),
    }
    if document_runtime is not None:
        env["EVOFLUX_DOCUMENT_RUNTIME_DIR"] = str(document_runtime)

    # Seed a minimal agents directory so the app's lifespan validation
    # (which requires exactly one agent with role: lead) passes.
    agents_dir = smoke_root / "config" / "agents"
    agents_dir.mkdir(parents=True, exist_ok=True)
    (agents_dir / "evoflux.md").write_text(
        "---\n"
        "name: evoflux\n"
        "role: lead\n"
        "model: __PROVIDER_MODEL__\n"
        "---\n"
        "# EvoFlux\n"
        "Smoke-test lead agent.\n"
    )

    # Use __main__.py path explicitly rather than ``-m app.cli`` so we
    # know *which* app.cli the interpreter finds — defends against a
    # vendored layout that buries app/ deeper later.
    cli_entry = site_packages / "app" / "cli" / "__main__.py"
    if not cli_entry.is_file():
        raise SystemExit(f"smoke test: missing CLI entry at {cli_entry}")

    bootstrap = (
        "import sys, runpy, site, faulthandler; "
        "faulthandler.dump_traceback_later(55, repeat=False); "
        "_site = sys.argv.pop(1); "
        "_entry = sys.argv.pop(1); "
        "site.addsitedir(_site); "
        "sys.argv[0] = _entry; "
        "runpy.run_path(_entry, run_name='__main__')"
    )

    smoke_cmd = [
        str(python_bin),
        "-c",
        bootstrap,
        str(site_packages),
        str(cli_entry),
        "serve",
        "--host",
        "127.0.0.1",
        "--port",
        "0",
        "--handshake",
        "--generate-token",
    ]
    print(">> " + " ".join(smoke_cmd))

    popen_kwargs: dict = {
        "stdout": subprocess.PIPE,
        "stderr": subprocess.PIPE,
        "env": env,
        "text": True,
    }
    if IS_WINDOWS:
        # CREATE_NEW_PROCESS_GROUP lets us terminate the child tree from
        # this process; start_new_session is Unix-only.
        popen_kwargs["creationflags"] = 0x0000_0200
    else:
        popen_kwargs["start_new_session"] = True

    proc = subprocess.Popen(smoke_cmd, **popen_kwargs)

    # Read output from background threads so the main thread can
    # enforce a real wall-clock timeout. ``subprocess.Popen.stdout`` is
    # buffered and blocking; without this scaffold, a child that goes
    # quiet hangs the smoke test indefinitely. Drain stderr too so
    # native/runtime warnings don't fill the pipe and so timeouts include
    # a useful tail for debugging.
    import queue as _queue
    import threading as _threading

    stdout_queue: "_queue.Queue[str | None]" = _queue.Queue()
    stdout_tail: list[str] = []
    stderr_tail: list[str] = []

    def _append_tail(buf: list[str], line: str, *, limit: int = 200) -> None:
        buf.append(line.rstrip())
        if len(buf) > limit:
            del buf[: len(buf) - limit]

    def _drain_stdout() -> None:
        assert proc.stdout is not None
        for line in iter(proc.stdout.readline, ""):
            _append_tail(stdout_tail, line)
            stdout_queue.put(line)
        stdout_queue.put(None)  # EOF sentinel

    def _drain_stderr() -> None:
        assert proc.stderr is not None
        for line in iter(proc.stderr.readline, ""):
            _append_tail(stderr_tail, line)

    stdout_reader = _threading.Thread(target=_drain_stdout, daemon=True)
    stderr_reader = _threading.Thread(target=_drain_stderr, daemon=True)
    stdout_reader.start()
    stderr_reader.start()

    def _timeout_message() -> str:
        out = "\n".join(stdout_tail[-80:]) or "<empty>"
        err = "\n".join(stderr_tail[-120:]) or "<empty>"
        return (
            "smoke test: handshake did not arrive in 60s\n"
            f"stdout tail:\n{out}\n"
            f"stderr tail:\n{err}"
        )

    payload: dict | None = None
    try:
        deadline = 60.0
        start = time.monotonic()
        while True:
            remaining = deadline - (time.monotonic() - start)
            if remaining <= 0:
                raise SystemExit(_timeout_message())
            try:
                line = stdout_queue.get(timeout=remaining)
            except _queue.Empty:
                raise SystemExit(_timeout_message())
            if line is None:
                err = "\n".join(stderr_tail)
                raise SystemExit(
                    f"smoke test: sidecar exited before handshake.\nstderr:\n{err[-4000:]}"
                )
            line = line.strip()
            if line.startswith("EVOFLUX_HANDSHAKE "):
                payload = json.loads(line.split(" ", 1)[1])
                break

        assert payload is not None
        port = payload["port"]
        token = payload["token"]
        base = f"http://127.0.0.1:{port}"
        print(f"smoke test: handshake ok: port={port} version={payload['version']}")

        # ── /api/health/live without token → must 401 ──────────────────────
        try:
            urllib.request.urlopen(f"{base}/api/health/live")
            # /api/health/live is exempt — no auth required even when token set.
            # That's intentional: orchestrator probes must work.
            print("smoke test: health/live reachable without token (exempt — expected)")
        except urllib.error.HTTPError as e:
            raise SystemExit(
                f"smoke test: health/live unexpectedly returned {e.code}"
            ) from e

        # ── /api/team/status without token → must 401 ──────────────────────
        try:
            urllib.request.urlopen(f"{base}/api/team/status", timeout=5)
            raise SystemExit(
                "smoke test: protected endpoint accepted request without token"
            )
        except urllib.error.HTTPError as e:
            if e.code != 401:
                raise SystemExit(
                    f"smoke test: expected 401 without token, got {e.code}"
                ) from e
            print("smoke test: protected endpoint correctly rejects missing token")

        # ── /api/team/status with token → must succeed (2xx or 503 OK) ─────
        req = urllib.request.Request(
            f"{base}/api/team/status",
            headers={"Authorization": f"Bearer {token}"},
        )
        try:
            urllib.request.urlopen(req, timeout=5)
            print("smoke test: protected endpoint accepts bearer token")
        except urllib.error.HTTPError as e:
            # 4xx other than 401 is OK (e.g. 404 if route changed), 5xx
            # is not — that signals the request reached the app but the
            # app blew up.
            if e.code == 401:
                raise SystemExit(
                    "smoke test: protected endpoint rejected valid token"
                ) from e
            if 500 <= e.code < 600:
                raise SystemExit(
                    f"smoke test: protected endpoint returned {e.code}"
                ) from e
            print(f"smoke test: protected endpoint returned {e.code} (acceptable)")

        # ── Bundled document runtime → manifest must resolve end-to-end ─────
        if document_runtime is not None:
            diagnostics_request = urllib.request.Request(
                f"{base}/api/diagnostics",
                headers={"Authorization": f"Bearer {token}"},
            )
            try:
                with urllib.request.urlopen(
                    diagnostics_request, timeout=10
                ) as response:
                    diagnostics_payload = json.loads(response.read().decode("utf-8"))
            except (
                urllib.error.HTTPError,
                urllib.error.URLError,
                json.JSONDecodeError,
            ) as exc:
                raise SystemExit(
                    f"smoke test: document runtime diagnostics failed: {exc}"
                ) from exc
            runtime_status = diagnostics_payload["runtime"]["document_runtime"]
            if runtime_status.get("available") is not True:
                raise SystemExit(
                    "smoke test: bundled document runtime is unavailable: "
                    f"{runtime_status.get('error', 'unknown error')}"
                )
            print(
                "smoke test: document runtime resolved: "
                f"{runtime_status.get('bundle_version')}"
            )

        # ── Browser-presence WebSocket → must complete a real upgrade ────────
        # Run the client with the bundled interpreter/site-packages too. This
        # catches the exact release regression where Uvicorn could serve HTTP
        # but the sidecar omitted every supported WebSocket transport.
        ws_url = (
            f"ws://127.0.0.1:{port}/api/team/sidecar-smoke/browser/presence"
            f"?_token={urllib.parse.quote(str(token), safe='')}"
        )
        ws_client = """
import asyncio, site, sys
site.addsitedir(sys.argv[1])
from websockets.asyncio.client import connect

async def check():
    async with connect(sys.argv[2], open_timeout=5, close_timeout=5) as websocket:
        await websocket.send("sidecar-smoke")

asyncio.run(check())
"""
        try:
            completed = subprocess.run(
                [
                    str(python_bin),
                    "-c",
                    ws_client,
                    str(site_packages),
                    ws_url,
                ],
                check=True,
                capture_output=True,
                env=env,
                text=True,
                timeout=15,
            )
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
            stdout = getattr(exc, "stdout", "") or "<empty>"
            stderr = getattr(exc, "stderr", "") or "<empty>"
            raise SystemExit(
                "smoke test: bundled WebSocket upgrade failed\n"
                f"client stdout:\n{stdout[-2000:]}\n"
                f"client stderr:\n{stderr[-4000:]}"
            ) from exc
        if completed.stderr:
            _append_tail(stderr_tail, completed.stderr)
        print("smoke test: bundled browser WebSocket upgrade succeeded")
    finally:
        if proc.poll() is None:
            if IS_WINDOWS:
                proc.terminate()
            else:
                proc.send_signal(signal.SIGTERM)
            try:
                proc.wait(timeout=5)
            except subprocess.TimeoutExpired:
                proc.kill()
        # Wipe the isolated smoke data dirs so we don't leak hundreds
        # of MB of throwaway state next to the bundle.
        shutil.rmtree(smoke_root, ignore_errors=True)


def human_bytes(n: int) -> str:
    for unit in ("B", "KB", "MB", "GB"):
        if n < 1024:
            return f"{n:.1f} {unit}"
        n //= 1024
    return f"{n:.1f} TB"


def report_size(root: Path, label: str) -> None:
    total = sum(p.stat().st_size for p in root.rglob("*") if p.is_file())
    print(f"  {label}: {human_bytes(total)}")


def main() -> int:
    ap = argparse.ArgumentParser(
        description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter
    )
    ap.add_argument(
        "--root", default=".", help="Project root containing pyproject.toml."
    )
    ap.add_argument(
        "--out", default="desktop/sidecar-bundle", help="Output bundle directory."
    )
    ap.add_argument(
        "--python-version",
        default="3.12",
        help="Major.minor Python version to bundle (default: 3.12).",
    )
    ap.add_argument(
        "--extras",
        default="",
        help="Comma-separated optional-dep extras to install (e.g. audio,azure-doc-intel,full).",
    )
    ap.add_argument(
        "--no-smoke",
        action="store_true",
        help="Skip the post-build smoke test (not recommended).",
    )
    ap.add_argument(
        "--no-zip-purelib",
        action="store_true",
        help=(
            "Keep pure-Python packages as loose files. By default Windows "
            "bundles safe packages into one zip to reduce Defender cold-start I/O."
        ),
    )
    ap.add_argument(
        "--document-runtime",
        default=os.environ.get(DOCUMENT_RUNTIME_SOURCE_ENV),
        help=(
            "Optional verified runtime directory/archive. Defaults to "
            f"${DOCUMENT_RUNTIME_SOURCE_ENV}; omitted from normal desktop builds."
        ),
    )
    ap.add_argument(
        "--document-runtime-sha256",
        default=os.environ.get(DOCUMENT_RUNTIME_SHA256_ENV),
        help=(
            "Required SHA-256 for an archive source; defaults to "
            f"${DOCUMENT_RUNTIME_SHA256_ENV}."
        ),
    )
    ap.add_argument(
        "--skip-document-runtime",
        action="store_true",
        help=(
            "Ignore a configured document runtime. Kept for compatibility; "
            "desktop builds no longer require a bundled runtime."
        ),
    )
    args = ap.parse_args()

    root = Path(args.root).resolve()
    out = Path(args.out).resolve()
    runtime_source: Path | None = None
    if args.document_runtime and not args.skip_document_runtime:
        runtime_source = Path(args.document_runtime).expanduser().resolve()
        if not runtime_source.exists():
            raise SystemExit(
                "configured document runtime source does not exist: "
                f"{runtime_source}"
            )
        if runtime_source == out or runtime_source.is_relative_to(out):
            raise SystemExit(
                "document runtime source must not be inside the output bundle"
            )

    if out.exists():
        print(f"removing existing {out}")
        shutil.rmtree(out)
    out.mkdir(parents=True)

    extras = [x.strip() for x in args.extras.split(",") if x.strip()]

    print(f"target python: {args.python_version}")
    print(f"target triple: {detect_target_triple()}")
    print(f"extras:        {extras or '(none — slim core)'}")
    print(f"output dir:    {out}")

    # ── 1. Fetch python-build-standalone ─────────────────────────────────
    install_root = out / "_python_install"
    uv_python_bin = fetch_python(args.python_version, install_root)
    python_target = out / "python"
    python_bin = normalise_python_dir(install_root, python_target, uv_python_bin)
    shutil.rmtree(install_root, ignore_errors=True)
    print(f"python binary: {python_bin}")

    # ── 2. Install EvoFlux + deps into site-packages ───────────────────
    site_packages = out / "site-packages"
    install_packages(python_bin, root, site_packages, extras)

    # ── 3. Stage the immutable document runtime ─────────────────────────
    document_runtime: Path | None = None
    if runtime_source is not None:
        document_runtime = out / "document-runtime"
        try:
            manifest = stage_document_runtime(
                runtime_source,
                document_runtime,
                expected_sha256=args.document_runtime_sha256,
                expected_platform=normalized_platform(),
                expected_architecture=normalized_architecture(),
            )
        except DocumentRuntimeError as exc:
            raise SystemExit(f"document runtime validation failed: {exc}") from exc
        print(
            "document runtime: "
            f"{manifest['bundle_version']} ({manifest['payload_sha256']})"
        )

    # ── 4. Strip caches/tests/etc. ──────────────────────────────────────
    saved = strip_bundle(site_packages)
    print(f"stripped: {human_bytes(saved)}")
    if IS_WINDOWS and not args.no_zip_purelib:
        packages_zipped, files_zipped = zip_pure_python_packages(site_packages)
        print(
            "zipimport: "
            f"packed {packages_zipped} pure-Python packages / {files_zipped} files"
        )

    # ── 5. Smoke test ───────────────────────────────────────────────────
    if not args.no_smoke:
        validate_migration_bundle(python_bin, site_packages)
        if document_runtime is not None:
            validate_document_runtime_bundle(
                python_bin, site_packages, document_runtime
            )
        smoke_test(python_bin, site_packages, document_runtime)

    # ── 6. Report ────────────────────────────────────────────────────────
    print("\n=== bundle summary ===")
    report_size(python_target, "python runtime")
    report_size(site_packages, "site-packages")
    if document_runtime is not None:
        report_size(document_runtime, "document runtime")
    report_size(out, "TOTAL")
    return 0


if __name__ == "__main__":
    sys.exit(main())
