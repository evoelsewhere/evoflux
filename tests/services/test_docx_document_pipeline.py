from __future__ import annotations

import json
from pathlib import Path
import zipfile

from docx import Document
from docx.oxml.ns import qn
import pytest

from app.services import docx_document_pipeline as pipeline
from app.services.office.runtime import file_sha256


def _new_project() -> pipeline.NewDocumentProject:
    return pipeline.NewDocumentProject.model_validate(
        {
            "schema_version": 1,
            "mode": "new",
            "title": "Decision Brief",
            "subtitle": "Editable Word-native document",
            "preset": "standard_business_brief",
            "header_pattern": "memo_masthead",
            "metadata": [{"label": "Status", "value": "Draft"}],
            "blocks": [
                {"type": "heading", "level": 1, "text": "Recommendation"},
                {"type": "paragraph", "text": "Approve the first wave."},
                {
                    "type": "table",
                    "columns": [
                        {"label": "Option", "width_dxa": 3000},
                        {"label": "Decision", "width_dxa": 6360},
                    ],
                    "rows": [["A", "Approve"]],
                },
            ],
        }
    )


def test_build_new_document_uses_native_word_objects(tmp_path: Path) -> None:
    output = tmp_path / "brief.docx"
    pipeline.build_new_document(_new_project(), output, tmp_path)
    document = Document(output)

    assert document.core_properties.title == "Decision Brief"
    assert any(paragraph.style.name == "Heading 1" for paragraph in document.paragraphs)
    assert len(document.tables) == 1
    assert document.tables[0].cell(1, 1).text == "Approve"
    assert document.sections[0].top_margin == pipeline.Inches(1)
    grid_widths = [
        column.get(qn("w:w"))
        for column in document.tables[0]._tbl.tblGrid.findall(qn("w:gridCol"))
    ]
    assert grid_widths == ["3000", "6360"]


def test_template_compose_preserves_unrelated_package_parts(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    source = tmp_path / "source.docx"
    pipeline.build_new_document(_new_project(), source, tmp_path)
    before = pipeline._package_hashes(source)
    project_path = tmp_path / "template-project.json"
    project_path.write_text(
        json.dumps(
            {
                "schema_version": 1,
                "mode": "template",
                "title": "Edit",
                "source_sha256": file_sha256(source),
                "template_confirmed": True,
                "edits": [
                    {
                        "operation": "replace_text",
                        "part": "word/document.xml",
                        "paragraph": 1,
                        "find": "Decision",
                        "replace": "Migration",
                    }
                ],
            }
        ),
        encoding="utf-8",
    )

    def fake_render(
        source: Path, render_dir: Path
    ) -> tuple[list[Path], list[dict[str, object]]]:
        render_dir.mkdir(parents=True, exist_ok=True)
        preview = render_dir / "page-001.png"
        preview.write_bytes(b"png")
        return [preview], []

    monkeypatch.setattr(pipeline, "render_docx_pages", fake_render)
    output = tmp_path / "edited.docx"
    result = pipeline.compose_document_project(
        project_path,
        source,
        output,
        asset_root=tmp_path,
        work_dir=tmp_path / "work",
    )

    assert result.passed
    assert result.metadata["reference_page_count"] == 1
    after = pipeline._package_hashes(output)
    assert before["word/styles.xml"] == after["word/styles.xml"]
    assert before["word/numbering.xml"] == after["word/numbering.xml"]
    assert before["word/document.xml"] != after["word/document.xml"]
    with zipfile.ZipFile(output) as package:
        assert b"Migration Brief" in package.read("word/document.xml")


def test_template_validation_detects_changed_source(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    pipeline.build_new_document(_new_project(), source, tmp_path)
    project = pipeline.TemplateDocumentProject.model_validate(
        {
            "schema_version": 1,
            "mode": "template",
            "title": "Edit",
            "source_sha256": file_sha256(source),
            "template_confirmed": True,
            "edits": [
                {
                    "operation": "replace_paragraph",
                    "paragraph": 0,
                    "text": "Updated",
                }
            ],
        }
    )
    source.write_bytes(b"changed")
    with pytest.raises(ValueError, match="changed after inspection"):
        pipeline.validate_document_project(project, source)
