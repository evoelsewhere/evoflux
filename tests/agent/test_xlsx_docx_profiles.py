"""Regression tests for profile-aware XLSX and DOCX generation."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook

from app.agent.builtin_skills.docx.scripts import qa as docx_qa
from app.agent.builtin_skills.docx.scripts import stylekit as docx_stylekit
from app.agent.builtin_skills.xlsx.scripts import qa as xlsx_qa
from app.agent.builtin_skills.xlsx.scripts import stylekit as xlsx_stylekit


def test_operational_xlsx_preserves_dense_editable_fields(tmp_path: Path) -> None:
    source = tmp_path / "operational.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Action Tracker"
    xlsx_stylekit.apply_workbook_profile(workbook, "operational")
    xlsx_stylekit.declare_content_contract(
        workbook,
        ["Action", "Owner", "Due Date", "Status", "Risk"],
    )
    xlsx_stylekit.style_title(
        sheet,
        "A1:F1",
        "90-day execution tracker",
        profile="operational",
    )
    headers = ["Action", "Owner", "Due Date", "Status", "Risk", "Notes"]
    for column, value in enumerate(headers, start=1):
        sheet.cell(4, column, value)
    xlsx_stylekit.style_header(sheet[4], profile="operational")
    for row in range(5, 17):
        sheet.cell(row, 1, f"Action {row - 4}")
        sheet.cell(row, 2, "Owner A")
        sheet.cell(row, 3, f"2026-08-{row - 3:02d}")
        sheet.cell(row, 4, "On track")
        sheet.cell(row, 5, "Low")
        sheet.cell(row, 6, "Editable operational note")
    xlsx_stylekit.style_body_range(
        sheet,
        "A5:F16",
        profile="operational",
    )
    xlsx_stylekit.set_column_widths(
        sheet,
        {"A": 24, "B": 14, "C": 13, "D": 13, "E": 10, "F": 28},
        profile="operational",
    )
    xlsx_stylekit.prepare_data_sheet(
        sheet,
        profile="operational",
        role="tracker",
        freeze="A5",
        auto_filter="A4:F16",
    )
    xlsx_stylekit.add_table(
        sheet,
        "A4:F16",
        name="ActionTracker",
        style="TableStyleMedium2",
    )
    xlsx_stylekit.add_list_validation(
        sheet,
        "D5:D200",
        ["Not started", "On track", "At risk", "Blocked", "Done"],
    )
    xlsx_stylekit.add_list_validation(
        sheet,
        "E5:E200",
        ["Low", "Medium", "High"],
    )
    xlsx_stylekit.add_status_formatting(sheet, "D5:D200")
    workbook.save(source)

    report = xlsx_qa.inspect_xlsx(source)

    assert report["errors"] == []
    assert report["warnings"] == []
    assert report["profile"] == "operational"
    assert report["required_fields"] == [
        "Action",
        "Owner",
        "Due Date",
        "Status",
        "Risk",
    ]
    metrics = report["sheet_metrics"][0]
    assert metrics["tables"] == 1
    assert metrics["validations"] == 2
    assert metrics["conditional_formats"] == 1
    assert metrics["freeze_panes"] == "A5"


def test_xlsx_content_contract_rejects_silently_dropped_field(
    tmp_path: Path,
) -> None:
    source = tmp_path / "missing-field.xlsx"
    workbook = Workbook()
    workbook.active["A1"] = "Action"
    xlsx_stylekit.apply_workbook_profile(workbook, "data-table")
    xlsx_stylekit.declare_content_contract(workbook, ["Action", "Approver"])
    workbook.save(source)

    report = xlsx_qa.inspect_xlsx(source)

    assert any("Approver" in error for error in report["errors"])


def _build_operational_docx(path: Path) -> None:
    document = Document()
    docx_stylekit.apply_theme(document, profile="operational-sop")
    docx_stylekit.declare_content_contract(
        document,
        ["Purpose", "Owner", "Procedure", "Acceptance"],
    )
    document.add_paragraph("Release readiness SOP", style="Title")
    document.add_paragraph("Purpose", style="Heading 1")
    document.add_paragraph(
        "Define a repeatable release gate with explicit ownership and evidence."
    )
    document.add_paragraph("Owner", style="Heading 2")
    document.add_paragraph("Release Manager")
    document.add_paragraph("Procedure", style="Heading 1")
    for step in (
        "Confirm the approved scope and dependency list.",
        "Validate automated checks and unresolved exceptions.",
        "Record the go/no-go decision and named approver.",
    ):
        document.add_paragraph(step, style="List Number")
    document.add_paragraph("Acceptance", style="Heading 1")
    table = document.add_table(rows=3, cols=3)
    values = [
        ["Check", "Owner", "Evidence"],
        ["Scope locked", "Product", "Approved backlog"],
        ["Quality gate", "Engineering", "Test report"],
    ]
    for row, row_values in zip(table.rows, values, strict=True):
        for cell, value in zip(row.cells, row_values, strict=True):
            cell.text = value
    docx_stylekit.style_table(
        table,
        widths_dxa=[3000, 1800, 4560],
        profile="operational-sop",
    )
    docx_stylekit.add_page_number_footer(document)
    document.save(str(path))


def test_operational_docx_uses_compact_profile_and_exact_table_geometry(
    tmp_path: Path,
) -> None:
    source = tmp_path / "operational.docx"
    _build_operational_docx(source)

    report = docx_qa.inspect_docx(source)

    assert report["errors"] == []
    assert report["warnings"] == []
    assert report["profile"] == "operational-sop"
    assert report["required_sections"] == [
        "Purpose",
        "Owner",
        "Procedure",
        "Acceptance",
    ]
    assert report["table_geometry"] == {
        "missing_geometry": 0,
        "fixed_rows": 0,
        "missing_repeating_headers": 0,
    }


def test_docx_qa_rejects_fixed_table_rows_that_can_clip(tmp_path: Path) -> None:
    source = tmp_path / "fixed-row.docx"
    document = Document()
    docx_stylekit.apply_theme(document, profile="compact-reference")
    document.add_paragraph("Checklist", style="Title")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Check"
    table.cell(0, 1).text = "Result"
    table.cell(1, 0).text = "Long editable check"
    table.cell(1, 1).text = "Pending"
    docx_stylekit.style_table(
        table,
        widths_dxa=[7000, 2360],
        profile="compact-reference",
    )
    row_properties = table.rows[1]._tr.get_or_add_trPr()
    height = OxmlElement("w:trHeight")
    height.set(qn("w:val"), "120")
    height.set(qn("w:hRule"), "exact")
    row_properties.append(height)
    document.save(str(source))

    report = docx_qa.inspect_docx(source)

    assert any("fixed height" in error for error in report["errors"])
